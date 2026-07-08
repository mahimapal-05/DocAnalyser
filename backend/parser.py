import io
from pypdf import PdfReader
import docx

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error parsing PDF: {str(e)}")

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error parsing Word Document: {str(e)}")

def parse_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception as e:
        raise ValueError(f"Error parsing text file: {str(e)}")

def extract_document_text(filename: str, file_bytes: bytes) -> str:
    """Router to parse document text based on file extension."""
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower_filename.endswith(".docx") or lower_filename.endswith(".doc"):
        return parse_docx(file_bytes)
    elif lower_filename.endswith((".txt", ".md", ".csv", ".json")):
        return parse_txt(file_bytes)
    else:
        # Fallback to text parsing
        try:
            return parse_txt(file_bytes)
        except Exception:
            raise ValueError(f"Unsupported file format: {filename}")
